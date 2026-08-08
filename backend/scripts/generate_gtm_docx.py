"""Generate an editable, shareable Word (.docx) version of the GTM Master Plan.

Output: docs/GTM-Strategy-ZenFragrances.docx
Charts rendered with matplotlib (strategy canvas, price ladder, courier,
tier mix, launch timeline) and embedded as images. Re-run after data changes.
"""

from __future__ import annotations

import io
import os
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
GREEN = RGBColor(0x05, 0x96, 0x69)
GREY = RGBColor(0x6B, 0x72, 0x80)

OUT_DIR = Path(__file__).resolve().parents[2] / "docs"
OUT_FILE = OUT_DIR / "GTM-Strategy-ZenFragrances.docx"


# ── Chart helpers ────────────────────────────────────────────────────────────

def _fig_bytes() -> io.BytesIO:
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close()
    buf.seek(0)
    return buf


def strategy_canvas() -> io.BytesIO:
    dims = ["Price(low)", "SKUs", "Web UX", "Stores", "CallCentre", "AgentPgm", "TeamComm", "ZeroAccess"]
    ffc = [8, 5, 6, 9, 8, 5, 1, 2]
    msc = [5, 9, 5, 8, 3, 8, 1, 2]
    zen = [3, 3, 4, 1, 1, 9, 9, 9]
    plt.figure(figsize=(7, 3.6))
    for data, label, color in ((ffc, "FFC", "#2563eb"), (msc, "M-Scents", "#16a34a"), (zen, "Zen", "#7c3aed")):
        plt.plot(dims, data, marker="o", label=label, color=color, linewidth=2)
    plt.ylim(0, 10)
    plt.title("Strategy Canvas — where we don't compete", fontsize=11, fontweight="bold")
    plt.legend(frameon=False)
    plt.grid(alpha=0.3)
    plt.tight_layout()
    return _fig_bytes()


def price_ladder() -> io.BytesIO:
    plt.figure(figsize=(5.5, 3.2))
    tiers = ["Budget", "Signature", "Premium"]
    wholesale = [40, 90, 175]
    retail = [70, 180, 350]
    x = range(len(tiers))
    plt.bar([i - 0.2 for i in x], wholesale, width=0.4, label="Wholesale", color="#a78bfa")
    plt.bar([i + 0.2 for i in x], retail, width=0.4, label="Retail (2×)", color="#7c3aed")
    plt.xticks(list(x), tiers)
    plt.ylabel("Rand / bottle")
    plt.title("3-Tier Price Ladder", fontsize=11, fontweight="bold")
    plt.legend(frameon=False)
    plt.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    return _fig_bytes()


def courier_economics() -> io.BytesIO:
    plt.figure(figsize=(5.5, 3.2))
    labels = ["1 budget", "1 mid", "3 mid", "5 mid", "10 mid", "20+"]
    values = [186, 68, 23, 14, 7, 0]
    bars = plt.bar(labels, values, color=["#dc2626", "#f59e0b", "#f59e0b", "#f59e0b", "#16a34a", "#16a34a"])
    plt.axhline(13, color="#7c3aed", linestyle="--", linewidth=1.5, label="Sustainability line (13%)")
    plt.ylabel("% of order value")
    plt.title("Courier cost vs order size", fontsize=11, fontweight="bold")
    plt.legend(frameon=False)
    plt.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    return _fig_bytes()


def tier_mix() -> io.BytesIO:
    plt.figure(figsize=(4.6, 3.4))
    plt.pie([60, 25, 15], labels=["Mid (engine)", "Premium (margin)", "Budget (acquisition)"],
            colors=["#7c3aed", "#a78bfa", "#ddd6fe"], startangle=90, autopct="%d%%",
            textprops={"fontsize": 9})
    plt.title("Tier mix target", fontsize=11, fontweight="bold")
    plt.tight_layout()
    return _fig_bytes()


def pre_launch_timeline() -> io.BytesIO:
    plt.figure(figsize=(7.4, 4.2))
    tasks = [
        ("Web store live (Vercel)", 0, 5, "#7c3aed"),
        ("WhatsApp live (Kapso)", 5, 4, "#7c3aed"),
        ("/flyer + price list live", 9, 3, "#7c3aed"),
        ("E2E order test", 12, 3, "#7c3aed"),
        ("Branding plan + colours", 0, 5, "#a78bfa"),
        ("Label + packaging design", 7, 10, "#a78bfa"),
        ("Product images (20 SKUs)", 7, 7, "#a78bfa"),
        ("Flyer restyle", 17, 3, "#a78bfa"),
        ("Finalize 20 SKUs + R30 lineup", 4, 5, "#16a34a"),
        ("Scent/gender mix confirmed", 9, 3, "#16a34a"),
        ("Order branded samples", 11, 5, "#2563eb"),
        ("Sniff + R30 price test", 16, 5, "#2563eb"),
        ("Refine SKU list", 21, 3, "#2563eb"),
        ("Recruit 50 seed agents", 14, 10, "#f59e0b"),
        ("Agent onboarding + first orders", 24, 5, "#f59e0b"),
        ("Recruit hawkers (Soweto/Tembisa)", 14, 10, "#f59e0b"),
        ("First hawker sales", 24, 5, "#f59e0b"),
        ("GO LIVE — Edition 1", 28, 2, "#dc2626"),
    ]
    y = range(len(tasks))
    for i, (name, start, dur, color) in enumerate(tasks):
        plt.barh(i, dur, left=start, color=color, edgecolor="white", height=0.6)
    plt.yticks(list(y), [t[0] for t in tasks], fontsize=8)
    plt.xlabel("Days from pre-launch start")
    plt.title("Pre-Launch (weeks −4 → 0)", fontsize=11, fontweight="bold")
    plt.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    return _fig_bytes()


def launch_timeline() -> io.BytesIO:
    plt.figure(figsize=(7.4, 3.4))
    tasks = [
        ("Validate mfg inventory", 0, 5, "#7c3aed"),
        ("Seed 20 SKUs + ladder", 5, 7, "#7c3aed"),
        ("Recruit 50 seed agents (GP)", 12, 21, "#16a34a"),
        ("Deploy web + flyer live", 12, 5, "#7c3aed"),
        ("FB Lead Ads test (R3k)", 22, 30, "#2563eb"),
        ("Budget tier via hawkers", 22, 30, "#16a34a"),
        ("Wholesalers tier open", 36, 21, "#2563eb"),
        ("Agent leaderboard + referrals", 29, 21, "#2563eb"),
        ("EC expansion (Mthatha/Mdantsane)", 52, 30, "#7c3aed"),
        ("First agent-run kiosk pilot", 67, 15, "#16a34a"),
    ]
    y = range(len(tasks))
    for i, (name, start, dur, color) in enumerate(tasks):
        plt.barh(i, dur, left=start, color=color, edgecolor="white", height=0.6)
    plt.yticks(list(y), [t[0] for t in tasks], fontsize=8)
    plt.xlabel("Days from start")
    plt.title("90-Day Launch Plan", fontsize=11, fontweight="bold")
    plt.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    return _fig_bytes()


# ── Document builder ─────────────────────────────────────────────────────────

def _style_table(table) -> None:
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def _add_chart(doc, buf: io.BytesIO) -> None:
    doc.add_picture(buf, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build() -> None:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Title page header
    title = doc.add_heading("Zen Fragrances — GTM Master Plan", level=0)
    for run in title.runs:
        run.font.color.rgb = PURPLE
    p = doc.add_paragraph("Go-To-Market Strategy · Consolidated 2026-08-08")
    p.runs[0].font.color.rgb = GREY
    doc.add_paragraph(
        "Zen Fragrances enables anyone in South Africa with a WhatsApp phone to own a perfume "
        "business — zero upfront cost, instant activation, two ways to earn. We eliminate barriers "
        "(no starter pack), reduce complexity (lean 20-SKU launch), raise agent autonomy (own pricing, "
        "team building, earnings transparency) and create new capabilities (WhatsApp-as-store, agent "
        "network as distribution). Our market isn't existing perfume buyers — it's the millions of "
        "South Africans looking for flexible income."
    )

    # 1. Strategy canvas
    doc.add_heading("1. Strategy Canvas — where we don't compete", level=1)
    _add_chart(doc, strategy_canvas())
    doc.add_paragraph(
        "Legend: FFC (national warehouses) · M-Scents (KZN stores) · Zen (people-network). "
        "We under-invest in what they compete on and own accessibility + team economics."
    )

    # 2. Channels
    doc.add_heading("2. Channel Strategy", level=1)
    _add_table(
        doc,
        ["Channel", "Verdict", "Capital", "Courier", "Role"],
        [
            ["WhatsApp (agents)", "✅ Core", "R0", "Bulk drop", "Ordering + communication (the moat)"],
            ["Web store", "✅ Keep", "R0", "Min R500", "Discovery, SEO, agent locator"],
            ["Agents", "✅ Core", "R0", "—", "The distribution army"],
            ["Wholesalers", "✅ Push", "R0", "Free-ship", "Bulk volume"],
            ["Hawkers", "✅ Budget tier", "R0", "Hand-sell", "Township impulse + penetration"],
            ["Stores (consignment)", "🟡 Defer", "R5–20K/store", "Bulk", "Only after profitable"],
        ],
    )

    # 3. Pricing
    doc.add_heading("3. Pricing — 3-Tier Ladder (20 launch SKUs)", level=1)
    _add_table(
        doc,
        ["Tier", "Sizes", "Wholesale", "Retail (2×)", "SKUs", "Role"],
        [
            ["🟢 Budget", "30ml", "R30–50", "R60–100", "6", "Acquisition / impulse / hand-sell"],
            ["🟡 Signature", "80–100ml", "R80–100", "R160–200", "10", "The revenue engine"],
            ["🔴 Premium/Oud", "100ml", "R150–200", "R300–400", "4", "Margin + gifting"],
            ["🎁 Add-ons", "sets/spritzers", "R8–45", "R20–90", "—", "Upsell, gift trigger"],
        ],
    )
    _add_chart(doc, price_ladder())
    doc.add_paragraph(
        "Budget-tier rule: never shipped alone (courier = 186% of value). It is a hand-sell "
        "acquisition tool, not a profit engine."
    )
    doc.add_paragraph(
        "INITIAL LAUNCH STRATEGY (2026-08-08): lead with the BUDGET tier at a starting price point "
        "of R30. Agents get 5% off this price (agent cost R28.50) and sell at ~2× (R60 retail, "
        "agent-set). Budget-first = lowest entry barrier, fastest agent recruitment, and hand-sell "
        "friendly (no courier on single budget bottles). Signature/premium tiers are added as agents grow."
    )

    # 4. Courier economics
    doc.add_heading("4. Courier Economics — the math that picks the channel", level=1)
    _add_table(
        doc,
        ["Order", "Value", "Courier", "% of value", "Viable?"],
        [
            ["1 budget", "R35", "R65", "186%", "❌ loss"],
            ["1 mid", "R95", "R65", "68%", "❌"],
            ["5 mid", "R475", "R65", "14%", "🟡"],
            ["10 mid", "R950", "R65", "7%", "✅"],
            ["20+ bottles", "R2,000+", "R0", "0%", "✅ best"],
        ],
    )
    _add_chart(doc, courier_economics())
    doc.add_paragraph(
        "Rule: courier sustainable at ≤10–13% of value → minimum shipped order ≈ R500 / 5 bottles. "
        "The agent is the last-mile courier (one bulk drop, local hand-sell)."
    )

    # 5. Distribution & payment
    doc.add_heading("5. Distribution — Penetration Ladder", level=1)
    _add_table(
        doc,
        ["Level", "What", "Capital", "Timing"],
        [
            ["L1 — Agent hubs", "Option 1: 1 hub per township, hawkers via agent code, Zen absorbs courier ≥R2k", "R0", "Month 1–3"],
            ["L2 — Buffer stock", "Hub agents hold 20–50 bottles, same-day local fulfillment", "R0", "Month 3–6"],
            ["L3 — Every town", "Every town has an agent; hub agents run depots", "R0–15K", "Month 6+"],
            ["L4 — Zen DCs", "Regional DCs (R50–150K each)", "R50–150K", "Year 2 (revenue-funded)"],
        ],
    )
    doc.add_heading("Payment flow — no agent money handling", level=2)
    doc.add_paragraph(
        "1) Hawker orders via WhatsApp (agent code) → 2) Zen sends Yoco/EFT/PayShap link → "
        "3) Hawker pays ZEN directly → 4) Zen ships one bulk drop to agent → 5) Agent distributes "
        "locally → 6) Zen pays agent 5% commission by bank transfer. Agents never touch money."
    )

    # 6. Geography
    doc.add_heading("6. Geography — where we enter", level=1)
    _add_table(
        doc,
        ["Priority", "Region", "Why", "Status"],
        [
            ["1", "Gauteng townships (Soweto, Tembisa, Mamelodi, Alexandra, Katlehong)", "FFC = city warehouses only; M-Scents absent", "Wave 1"],
            ["2", "Eastern Cape (Mthatha, Mdantsane, Gqeberha)", "Both competitors thin; M-Scents Mthatha signal", "Wave 2"],
            ["⛔", "KZN (Durban, PMB)", "Saturated — M-Scents 14 stores + FFC", "Avoid"],
        ],
    )
    doc.add_paragraph("Only competitor overlap is KZN. Everywhere else is contested by one player or none.")

    # 7. Pre-launch + launch plan
    doc.add_heading("7. Pre-Launch + 90-Day Launch Plan", level=1)
    _add_chart(doc, pre_launch_timeline())
    _add_table(
        doc,
        ["#", "Track", "Deliverable", "Week", "Exit criteria (Gate)"],
        [
            ["1", "Site & WA ready", "Web store + WhatsApp + /flyer + price list live", "−4→−3", "E2E order test passes + UX audit (checklist.design)"],
            ["2", "Branding plan", "Colours + logo + brand kit (tasteskill brandkit)", "−4", "Founder sign-off"],
            ["3", "Branding", "Label, packaging, product images, flyer restyle; no-ai-slop copy pass", "−3→−2", "20 SKUs have branded images; copy sounds human"],
            ["4", "Fast movers", "Finalize 20 SKUs (6/10/4) + R30 lineup", "−3", "SKU list frozen"],
            ["5", "Samples / market test", "Branded samples; sniff + R30 price test", "−2→−1", "SKU list refined from feedback"],
            ["6", "Activate agents", "50 seed agents; onboarding + first orders", "−1→0", "50 active · ≥10 pilot orders"],
            ["7", "Activate hawkers", "Hawkers (Soweto/Tembisa); first taxi-rank sales", "−1→0", "First hawker sales recorded"],
            ["8", "GO LIVE", "Edition-1 flyer + broadcasts; transitions.dev polish", "0", "Live"],
        ],
    )
    _add_chart(doc, launch_timeline())

    # 8. KPIs
    doc.add_heading("8. KPI Dashboard (North Star)", level=1)
    _add_table(
        doc,
        ["Metric", "M1", "M3", "M6"],
        [
            ["Active agents", "100", "500", "2,000"],
            ["Monthly orders", "200", "1,500", "6,000"],
            ["Monthly GMV", "R70K", "R600K", "R2.7M"],
            ["Courier % of revenue", "< 8%", "< 6%", "< 5%"],
            ["Budget tier share", "< 60%", "< 40%", "< 20%"],
            ["Avg agent order", "R1,500", "R1,800", "R2,000+"],
            ["CAC / active agent", "< R80", "< R60", "< R50"],
        ],
    )
    _add_chart(doc, tier_mix())

    # 9. GTM approaches considered
    doc.add_heading("9. GTM Approaches Considered (strengths & weaknesses)", level=1)
    _add_table(
        doc,
        ["Approach", "Strengths", "Weaknesses", "Verdict"],
        [
            ["WhatsApp-first agent platform", "Lowest friction in SA; no app/login; network effects; low-literacy friendly", "No physical sniff test; needs agent recruitment; courier forces bulk", "✅ Core"],
            ["National warehouse + web + call centre (FFC)", "National reach; proven; brand recognition", "R960 barrier; web-only friction; high capex; shallow in communities", "❌ Not ours"],
            ["Physical retail stores (M-Scents)", "Community trust; sniff test; cash sales", "Capital-heavy; regional; rent/staff; slow scale", "🟡 Defer — agent hubs replace stores"],
            ["Direct-to-consumer web", "Full margin; SEO", "Single-bottle courier kills economics; needs ad spend", "⚠️ Weakest channel"],
            ["Branded distribution (license Motala/P2D)", "Instant brand trust; gift sets; price ladder", "Thin margin; no brand equity for US; replaceable middleman", "🟡 Later, optional licensing"],
            ["Own-brand manufacturing (private label)", "Full margin; full brand equity; owns the vertical; not replaceable", "Needs own packaging/brand investment; brand trust builds over time", "✅ Chosen (day 1)"],
            ["Hawker / micro-retail network", "Zero courier; impulse; township penetration", "Small per-unit margin; management heavy", "✅ Budget-tier channel"],
        ],
    )
    doc.add_paragraph(
        "Why the AGENT route initially: (1) courier economics make single-bottle direct-to-consumer "
        "unviable (R65 courier on a R95 bottle); (2) retail stores need capital we don't have; "
        "(3) wholesaling makes us a replaceable middleman; (4) agents give zero-capital, WhatsApp-native, "
        "network-effect distribution AND build the army that later supports own-brand retail. "
        "Fastest path to first revenue while defending our moat."
    )

    # 10. SWOT
    doc.add_heading("10. SWOT (readable summary)", level=1)
    _add_table(
        doc,
        ["💪 Strengths (internal)", "🩸 Weaknesses (internal)"],
        [
            ["Zero barrier (no R960 starter pack)", "No physical presence / no sniff test"],
            ["WhatsApp-first (95% of SA have it)", "No brand recognition yet — own-brand must be built"],
            ["Two-sided earning (margin + 5% team)", "Single manufacturer (Focus Logic) — dual-source mandatory"],
            ["6 role dashboards + agent locator", "No card payments on WhatsApp yet"],
            ["Confirmation-before-cart; click-through links", "0 of 99 SKUs seeded; single-size plan"],
            ["Automated mfg forwarding; lean iteration", "Budget-first means thin per-unit margin initially"],
        ],
    )
    _add_table(
        doc,
        ["🚀 Opportunities (external)", "⚠️ Threats (external)"],
        [
            ["R750B township economy + side-hustle culture", "FFC adds WhatsApp / drops starter price"],
            ["WhatsApp = SA super-app (join habit, not build)", "Copycat entrants (speed is the moat)"],
            ["FFC's R960 barrier = our lead source", "WhatsApp policy changes (Kapso buffers)"],
            ["Social commerce boom → agents = free sales force", "Manufacturer failure (dual-source by mth 3)"],
            ["Agent network = last-mile distribution (no rent)", "Pyramid perception (commission on orders only)"],
            ["Corporate gifting; SADC; data monetization", "Courier theft; counterfeit stigma; rand volatility"],
        ],
    )

    # 11. Competitor landscape
    doc.add_heading("11. Competitor Landscape (factual)", level=1)
    _add_table(
        doc,
        ["Competitor", "SKUs", "Sizes", "Wholesale", "Retail", "Starter", "Footprint", "Ordering"],
        [
            ["FFC", "42", "30ml only", "R19–94", "R40–190", "R960", "15 warehouses (national) + kiosks", "Web + call centre"],
            ["M-Scents", "241", "15ml–200ml", "R15–200", "R30–400+", "R800–3,750", "21 stores (ALL KZN)", "Web only"],
            ["Fragrance Boutique", "40+", "50/100ml", "—", "R219–419", "Application", "1 store (CPT)", "Web + WhatsApp + physical"],
            ["Perfumes for Africa", "40+", "5ml–100ml", "—", "R12–133", "—", "1 store (CPT)", "Web + WhatsApp + physical"],
            ["SensoryFX / Sensetek", "B2B", "—", "—", "—", "—", "Centurion / Sandton", "—"],
        ],
    )
    doc.add_paragraph(
        "Facts that drive the plan: only competitor overlap is KZN (Durban + PMB); Gauteng townships have "
        "NO competitor community presence; both leaders (FFC, M-Scents) are web-only with high entry "
        "barriers; no competitor has team commissions, WhatsApp-first ordering, or agent-as-distribution."
    )

    # 12. Owning the vertical
    doc.add_heading("12. Owning the Vertical — Own-Brand First", level=1)
    doc.add_paragraph(
        "DECISION (2026-08-08): we PRODUCE and SELL perfumes under our OWN brand (Zen) from day 1, via "
        "contract manufacturer. Licensing third-party brands (Motala/P2D/Parfumo) is a LATER, optional "
        "expansion — not the launch path."
    )
    _add_table(
        doc,
        ["Assumption", "Impact of own-brand-first"],
        [
            ["Moat", "Stronger — we're a platform + a brand, not a replaceable distributor"],
            ["Brand", "Critical from day 1 — build Zen equity ourselves (packaging, quality, brand colours). No borrowed trust"],
            ["Margin ceiling", "Higher from day 1 — full margin, no distributor/licensing cut"],
            ["Manufacturer dependency", "Critical — contract manufacturer makes our oil. Dual-source (SensoryFX/Sensetek) mandatory"],
            ["Zero-capital launch", "Mostly unchanged — contract manufacturer holds inventory; own-label packaging is the launch cost"],
            ["Counterfeit perception", "More relevant — our own brand must not look fake; quality packaging non-negotiable"],
        ],
    )

    # 13. Payment & collection
    doc.add_heading("13. Payment Pathways & Collection / Courier", level=1)
    doc.add_paragraph("Payment pathways (hawker/agent pays ZEN directly — agents never handle money):")
    _add_table(
        doc,
        ["Payment method", "How it works", "Clears"],
        [
            ["Yoco card link", "Payment link sent on WhatsApp / web checkout", "Instant"],
            ["EFT + POP", "Bank transfer, then send proof-of-payment image", "1–2 working days (manual verify)"],
            ["PayShap / mobile money", "Instant bank-to-bank via phone", "Instant"],
            ["Cash (exception)", "Cash-only hawkers — collected by agent, capped per agent", "Paid by agent to Zen (exception only)"],
        ],
    )
    doc.add_paragraph(
        "Collection / courier once payment is made:\n"
        "1) Payment confirmed (POP verified / idempotency)\n"
        "2) Zen consolidates the agent's group of paid orders\n"
        "3) Courier Guy to agent (R65 flat; FREE over R2,000) OR Pargo/Pudo pickup point (rural — agent collects)\n"
        "4) Agent distributes locally (hand-sell)\n"
        "Rule: budget bottles are hand-sold or bundled — never couriered alone."
    )

    # 14. Positioning + 15. Open items
    doc.add_heading("14. Competitive Positioning", level=1)
    _add_table(
        doc,
        ["vs", "Our line"],
        [
            ["FFC", "Start with 1 bottle, R65 — vs their R960 starter pack + web-only"],
            ["M-Scents", "No starter pack, WhatsApp-first, team commissions, Gauteng — vs their R800–R3,750 + KZN-only"],
            ["Fragrance Boutique", "Become an agent instantly — vs application-based"],
        ],
    )

    doc.add_heading("15. Open Items", level=1)
    for item in [
        "Verify M-Scents Mthatha store (not on their locator yet)",
        "Set FLYER_WHATSAPP in Railway for the flyer CTA",
        "Confirm Focus Logic can private-label (our Zen brand) + holds inventory (the model depends on it)",
        "Define Zen own-brand identity: name/bottle/packaging/brand colours (launch-critical)",
        "Confirm dual-source private-label capability (SensoryFX/Sensetek) — mandatory hedge",
        "Consider licensing Motala/P2D/Parfumo LATER (optional expansion, not launch)",
        "Confirm R30 budget price point + 5% agent discount mechanics",
        "Seed 20 SKUs under the Zen brand with thumbnail_url images",
        "Refine flyer styling once brand colours are chosen",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph("Zen Fragrances — GTM Master Plan · generated 2026-08-08 · editable & shareable")
    footer.runs[0].font.color.rgb = GREY
    footer.runs[0].font.size = Pt(8)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_FILE))
    print(f"wrote {OUT_FILE} ({OUT_FILE.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    build()
